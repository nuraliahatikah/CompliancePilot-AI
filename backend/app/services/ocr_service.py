import io
import logging
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from PIL import Image

from app.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class OCRService:
    """Extract text from documents with optional Tesseract OCR for scanned content."""

    SUPPORTED_IMAGE_TYPES = {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}
    SUPPORTED_DOC_TYPES = {".pdf", ".docx", ".txt"}

    def __init__(self) -> None:
        if settings.tesseract_cmd:
            import pytesseract

            pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

    def extract_text(self, file_path: str, file_type: str) -> tuple[str, bool]:
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore"), False

        if extension == ".docx":
            return self._extract_docx(path), False

        if extension == ".pdf":
            return self._extract_pdf(path)

        if extension in self.SUPPORTED_IMAGE_TYPES:
            return self._ocr_image(path), True

        raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_docx(self, path: Path) -> str:
        doc = DocxDocument(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n\n".join(paragraphs)

    def _extract_pdf(self, path: Path) -> tuple[str, bool]:
        doc = fitz.open(str(path))
        text_parts: list[str] = []
        ocr_needed = False

        for page in doc:
            page_text = page.get_text("text").strip()
            if page_text:
                text_parts.append(page_text)
            else:
                ocr_needed = True
                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = self._run_tesseract(img)
                if ocr_text.strip():
                    text_parts.append(ocr_text)

        doc.close()
        combined = "\n\n".join(text_parts)
        if not combined.strip() and not ocr_needed:
            ocr_needed = True
            combined = self._ocr_pdf_pages(path)
        return combined, ocr_needed

    def _ocr_pdf_pages(self, path: Path) -> str:
        doc = fitz.open(str(path))
        text_parts: list[str] = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text_parts.append(self._run_tesseract(img))
        doc.close()
        return "\n\n".join(text_parts)

    def _ocr_image(self, path: Path) -> str:
        img = Image.open(str(path))
        return self._run_tesseract(img)

    def _run_tesseract(self, image: Image.Image) -> str:
        try:
            import pytesseract

            return pytesseract.image_to_string(image, lang="eng")
        except Exception as exc:
            logger.warning("Tesseract OCR unavailable: %s", exc)
            return ""


ocr_service = OCRService()
